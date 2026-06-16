"""source_loader.py - Dual-entry ingestion for the unified Video.AI pipeline.

Accepts 5 input types + pasted text and produces a uniform `SourceDocument`:

  - .txt   : UTF-8, strip BOM, normalize line endings
  - .md    : same as .txt + strip YAML front-matter into metadata
  - .pdf   : pypdf text extraction (no OCR; scanned PDFs raise with a clear error)
  - .docx  : python-docx, paragraphs in order, headings preserved as metadata
  - URL    : trafilatura main-content extraction + requests for fetching
             (sets a User-Agent header per Wikimedia ToS)
  - paste  : string passed through directly

Detection: pathlib.Path.suffix for files, http(s):// prefix for URLs,
string type for paste.

Pure functions - no LLM calls, no GPU.
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

log = logging.getLogger(__name__)


SUPPORTED_EXTENSIONS = (".txt", ".md", ".pdf", ".docx")


class SourceLoaderError(Exception):
    """Raised when a source cannot be loaded (unsupported type, missing file, etc.)."""


@dataclass
class SourceDocument:
    """Uniform output of load_source()."""

    text: str
    word_count: int
    language: str
    source_type: str
    metadata: dict[str, Any] = field(default_factory=dict)


def _word_count(text: str) -> int:
    return len(text.split())


def _detect_language(text: str) -> str:
    total_alpha = sum(1 for c in text if c.isalpha())
    if total_alpha == 0:
        return "unknown"
    deva = sum(1 for c in text if "\u0900" <= c <= "\u097f")
    ratio = deva / total_alpha
    if ratio >= 0.5:
        return "hi"
    if ratio >= 0.1:
        return "mixed"
    return "en"


def _check_oversize(text: str, max_words: int) -> None:
    wc = _word_count(text)
    if wc > max_words:
        log.warning(
            f"source_loader: text is {wc} words, exceeds soft cap of {max_words}. Proceeding."
        )


def _normalize_text(raw: str) -> str:
    """Strip BOM, normalize CRLF/CR -> LF."""
    if raw.startswith("\ufeff"):
        raw = raw[1:]
    raw = raw.replace("\r\n", "\n").replace("\r", "\n")
    return raw


def _strip_md_frontmatter(text: str) -> tuple[str, dict[str, Any]]:
    """Strip YAML front-matter delimited by --- at the start of the file.

    Returns (body_text, metadata_dict). Empty metadata when no front-matter.
    """
    if not text.startswith("---"):
        return text, {}
    lines = text.split("\n")
    if not lines or lines[0].strip() != "---":
        return text, {}
    end_idx = None
    for i in range(1, len(lines)):
        if lines[i].strip() == "---":
            end_idx = i
            break
    if end_idx is None:
        return text, {}
    front_lines = lines[1:end_idx]
    body = "\n".join(lines[end_idx + 1 :]).lstrip("\n")
    metadata: dict[str, Any] = {}
    for line in front_lines:
        if ":" in line and not line.strip().startswith("#"):
            key, _, value = line.partition(":")
            metadata[key.strip()] = value.strip().strip('"').strip("'")
    return body, metadata


def _load_txt(path: Path) -> SourceDocument:
    raw = path.read_text(encoding="utf-8", errors="replace")
    text = _normalize_text(raw)
    return SourceDocument(
        text=text,
        word_count=_word_count(text),
        language=_detect_language(text),
        source_type="txt",
        metadata={"path": str(path.resolve()), "size_bytes": path.stat().st_size},
    )


def _load_md(path: Path) -> SourceDocument:
    raw = path.read_text(encoding="utf-8", errors="replace")
    text = _normalize_text(raw)
    body, front_meta = _strip_md_frontmatter(text)
    return SourceDocument(
        text=body,
        word_count=_word_count(body),
        language=_detect_language(body),
        source_type="md",
        metadata={
            "path": str(path.resolve()),
            "size_bytes": path.stat().st_size,
            "front_matter": front_meta,
        },
    )


def _load_pdf(path: Path) -> SourceDocument:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise SourceLoaderError(
            "pypdf is required to load .pdf files. Install with: pip install pypdf>=4.0"
        ) from e

    reader = PdfReader(str(path))
    page_texts: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            extracted = page.extract_text() or ""
        except Exception as e:
            log.warning(f"source_loader: PDF page {i} extraction failed: {e}")
            extracted = ""
        page_texts.append(extracted)
    text = _normalize_text("\n".join(page_texts))
    if not text.strip():
        raise SourceLoaderError(
            f"PDF '{path}' contains no extractable text (likely scanned/no OCR). "
            "OCR is not supported - supply a .txt or .md version."
        )
    return SourceDocument(
        text=text,
        word_count=_word_count(text),
        language=_detect_language(text),
        source_type="pdf",
        metadata={
            "path": str(path.resolve()),
            "size_bytes": path.stat().st_size,
            "page_count": len(reader.pages),
        },
    )


def _load_docx(path: Path) -> SourceDocument:
    try:
        from docx import Document
    except ImportError as e:
        raise SourceLoaderError(
            "python-docx is required to load .docx files. "
            "Install with: pip install python-docx>=1.1"
        ) from e

    doc = Document(str(path))
    paragraph_data: list[dict[str, str]] = []
    text_parts: list[str] = []
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "Normal"
        is_heading = style_name.startswith("Heading")
        paragraph_data.append({"style": style_name, "text": para.text, "is_heading": is_heading})
        text_parts.append(para.text)
    text = _normalize_text("\n".join(text_parts))
    headings = [
        {"style": p["style"], "text": p["text"]}
        for p in paragraph_data
        if p["is_heading"] and p["text"].strip()
    ]
    core_props = doc.core_properties
    author = (core_props.author or "") if core_props else ""
    return SourceDocument(
        text=text,
        word_count=_word_count(text),
        language=_detect_language(text),
        source_type="docx",
        metadata={
            "path": str(path.resolve()),
            "size_bytes": path.stat().st_size,
            "headings": headings,
            "author": author,
            "paragraph_count": len(paragraph_data),
        },
    )


def _load_url(url: str, config: dict | None) -> SourceDocument:
    import requests

    try:
        import trafilatura
    except ImportError as e:
        raise SourceLoaderError(
            "trafilatura is required to load URLs. Install with: pip install trafilatura>=1.6"
        ) from e

    from utils.url_security import validate_source_url

    cfg_source = (config or {}).get("source", {})
    user_agent = cfg_source.get("user_agent", "VideoAI/6.0 (+https://github.com/...)")
    timeout_s = int(cfg_source.get("url_timeout_s", 30))

    # H6 fix: cap the fetch size so a huge/malicious URL cannot exhaust
    # memory via resp.text. Declared size is checked first; the actual body
    # size is re-checked after download (covers chunked responses).
    max_bytes = int(cfg_source.get("max_fetch_bytes", 10 * 1024 * 1024))

    # SSRF: reject private/metadata/link-local IPs, localhost, file://, etc.
    validated_url = validate_source_url(url)

    log.info(f"source_loader: fetching {validated_url} (timeout={timeout_s}s)")
    try:
        resp = requests.get(validated_url, headers={"User-Agent": user_agent}, timeout=timeout_s, allow_redirects=False)
        resp.raise_for_status()
    except requests.RequestException as e:
        raise SourceLoaderError(f"URL fetch failed: {e}") from e

    # Follow redirects manually, validating each hop
    from urllib.parse import urljoin

    max_redirects = 5
    redirect_count = 0
    current_url = validated_url
    while resp.is_redirect and redirect_count < max_redirects:
        next_url = resp.headers.get("Location", "")
        if not next_url:
            break
        # Resolve relative redirects against the current URL
        next_url = urljoin(current_url, next_url)
        validated_next = validate_source_url(next_url)
        current_url = validated_next
        redirect_count += 1
        resp = requests.get(validated_next, headers={"User-Agent": user_agent}, timeout=timeout_s, allow_redirects=False)
        resp.raise_for_status()

    try:
        declared = int(resp.headers.get("Content-Length", 0))
    except (TypeError, ValueError):
        declared = 0
    if declared > max_bytes:
        raise SourceLoaderError(
            f"URL response too large: {declared} bytes (cap: {max_bytes}). "
            "Raise source.max_fetch_bytes in config to allow."
        )

    raw_html = resp.text
    if len(raw_html) > max_bytes:
        raise SourceLoaderError(
            f"URL response too large: {len(raw_html)} chars (cap: {max_bytes})."
        )
    extracted = trafilatura.extract(raw_html, include_comments=False, include_tables=False) or ""
    text = _normalize_text(extracted)
    if not text.strip():
        raise SourceLoaderError(
            f"URL '{url}' yielded no main-content text. Site may be JS-rendered or paywalled."
        )
    metadata: dict[str, Any] = {
        "url": url,
        "fetch_date": __import__("datetime")
        .datetime.now(__import__("datetime").timezone.utc)
        .isoformat(),
        "status_code": resp.status_code,
        "raw_html_size": len(raw_html),
        "extracted_size": len(text),
    }
    if resp.url and resp.url != url:
        metadata["final_url"] = resp.url
    return SourceDocument(
        text=text,
        word_count=_word_count(text),
        language=_detect_language(text),
        source_type="url",
        metadata=metadata,
    )


def _load_paste(text: str) -> SourceDocument:
    text = _normalize_text(text)
    return SourceDocument(
        text=text,
        word_count=_word_count(text),
        language=_detect_language(text),
        source_type="paste",
        metadata={"char_count": len(text)},
    )


def _is_url(s: str) -> bool:
    return bool(re.match(r"^https?://", s.strip(), re.IGNORECASE))


def _validate_extension(path: Path, allowed: tuple[str, ...]) -> None:
    if path.suffix.lower() not in allowed:
        raise SourceLoaderError(
            f"Unsupported file extension '{path.suffix}'. Allowed: {', '.join(allowed)}"
        )


def load_source(source: Any, config: dict | None = None) -> SourceDocument:
    """Load a source from one of 5 input types + pasted text.

    Args:
        source: One of:
            - pathlib.Path / str ending in .txt|.md|.pdf|.docx
            - str starting with http:// or https://  (URL)
            - str (anything else) treated as pasted text
        config: Optional config dict; reads `source.allowed_extensions`,
            `source.max_words`, `source.url_timeout_s`, `source.user_agent`.

    Returns:
        SourceDocument with text, word_count, language, source_type, metadata.

    Raises:
        SourceLoaderError on unsupported type, missing file, or fetch failure.
    """
    cfg_source = (config or {}).get("source", {})
    allowed = tuple(cfg_source.get("allowed_extensions", list(SUPPORTED_EXTENSIONS)))
    max_words = int(cfg_source.get("max_words", 50000))

    if isinstance(source, Path):
        path = source
    elif isinstance(source, str) and not _is_url(source):
        stripped = source.strip()
        # H7 fix: only treat the string as a file path when it could actually
        # be one - single line and bounded length. Pasted documents contain
        # newlines; without this check, paste text accidentally ending in an
        # allowed extension was read from disk (arbitrary local file read).
        looks_like_path = (
            "\n" not in stripped
            and len(stripped) <= 4096
            and any(stripped.lower().endswith(ext) for ext in allowed)
        )
        if looks_like_path:
            path = Path(stripped)
        else:
            doc = _load_paste(source)
            _check_oversize(doc.text, max_words)
            return doc
    elif isinstance(source, str) and _is_url(source):
        doc = _load_url(source.strip(), config)
        _check_oversize(doc.text, max_words)
        return doc
    else:
        raise SourceLoaderError(
            f"Unsupported source type: {type(source).__name__}. "
            "Expected Path, str (file path / URL / pasted text), or None."
        )

    if not path.exists():
        raise SourceLoaderError(f"File not found: {path}")
    _validate_extension(path, allowed)

    ext = path.suffix.lower()
    if ext == ".txt":
        doc = _load_txt(path)
    elif ext == ".md":
        doc = _load_md(path)
    elif ext == ".pdf":
        doc = _load_pdf(path)
    elif ext == ".docx":
        doc = _load_docx(path)
    else:
        raise SourceLoaderError(f"Unhandled extension: {ext}")

    _check_oversize(doc.text, max_words)
    return doc
